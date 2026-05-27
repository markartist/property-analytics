from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/mark/Property_Analytics")
ARCHIVE_DIR = ROOT / "docs" / "archive" / "research" / "andrew_web_platform_architecture_2026-05-26"
SOURCE = ARCHIVE_DIR / "WEB_PLATFORM_ARCHITECTURE_FOR_ANDREW_2026-05-26.md"
OUT = ARCHIVE_DIR / "Venterra_Web_Platform_Architecture_for_Andrew_2026-05-26.docx"

BLUE = "1F4D78"
ACCENT = "2E74B5"
DARK = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "000000"
BORDER = "D0D7DE"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size="6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(w))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def paragraph_bottom_border(paragraph, color=ACCENT, size="12", space="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def clean_inline(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text)


def add_inline_runs(paragraph, text: str, size=11, color=BLACK):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=size, color=DARK, name="Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        else:
            run = paragraph.add_run(part.replace("`", ""))
            set_run_font(run, size=size, color=color)


def add_para(doc, text="", style=None, before=0, after=6, line=1.10):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        add_inline_runs(p, text)
    return p


def add_callout(doc, title: str, body: str, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="C9D6E2", size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    add_inline_runs(p2, body, size=10.5, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    set_cell_border(cell, color="D8DEE4", size="6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.rstrip().splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9.5, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        rows.pop(1)
    return rows


def add_markdown_table(doc, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    header = [cell.lower() for cell in rows[0]]
    if cols == 2 and header == ["public venterra path", "kinsta-hosted resi site"]:
        widths = [3900, 5460]
    elif cols == 2:
        widths = [2600, 6760]
    elif cols == 3:
        widths = [2200, 2600, 4560]
    else:
        base = 9360 // cols
        widths = [base] * cols
        widths[-1] += 9360 - sum(widths)
    set_table_width(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            text = row[c_idx] if c_idx < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline_runs(p, text, size=9.5 if len(text) > 55 else 10.2)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK)
            else:
                set_cell_shading(cell, WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_process_flow(doc):
    rows = [
        ["Step", "System action"],
        ["1", "User requests a canonical property URL on venterraliving.com."],
        ["2", "Kinsta receives the HTTPS request and resolves the city/state property slug."],
        ["3", "Kinsta direct reverse-proxy routing sends the request to the correct Resi WordPress property site."],
        ["4", "The Resi property site returns the page, media references, metadata, and leasing calls to action."],
        ["5", "Kinsta returns the response under the Venterra URL, preserving the subdirectory structure and trailing slash."],
        ["6", "Tours, applications, and quotes securely hand off to online.venterraliving.com with approved attribution ID variables."],
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [900, 8460])
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            add_inline_runs(p, text, size=10.2)
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK)
            else:
                set_cell_shading(cell, WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_cover(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Venterra Web Platform Architecture")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Internal Technical Review | Page ")
    set_run_font(run, size=9, color=MUTED)
    add_field(footer, "PAGE")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("INTERNAL TECHNICAL REVIEW")
    set_run_font(run, size=10, color=ACCENT, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Venterra Web Platform Architecture")
    set_run_font(run, size=28, color=DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Kinsta-hosted public web framework for Resi property routing, SSL, canonical URLs, secure portal handoff, and IT governance")
    set_run_font(run, size=13.5, color=MUTED)
    paragraph_bottom_border(p, color=ACCENT, size="16", space="10")

    meta = doc.add_table(rows=5, cols=2)
    set_table_width(meta, [1900, 7460])
    items = [
        ("Status", "Proposed framework"),
        ("Date", "05/26/2026"),
        ("Audience", "Andrew / IT Leadership"),
        ("Scope", "Public Venterra web platform from SSL to Kinsta hosting, Resi property routing, and end-user experience"),
        ("Canonical model", "https://venterraliving.com/apartments/{property-slug}-{city}-{state}/"),
    ]
    for i, (label, value) in enumerate(items):
        for j, txt in enumerate((label, value)):
            cell = meta.cell(i, j)
            cell.text = ""
            set_cell_shading(cell, LIGHT_BLUE if j == 0 else WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, txt, size=10.5)
            for r in p.runs:
                if j == 0:
                    r.bold = True
                    r.font.color.rgb = RGBColor.from_string(DARK)

    add_callout(
        doc,
        "Architecture Summary",
        "The public web platform will be hosted on Kinsta, use venterraliving.com as the SSL-secured non-www canonical domain, preserve /apartments/{property-slug}-{city}-{state}/ URLs with trailing slashes, and securely hand off tours, applications, and quotes to online.venterraliving.com with approved attribution IDs.",
        fill="EEF6FF",
    )
    doc.add_page_break()


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, ACCENT, 16, 8),
        ("Heading 2", 13, ACCENT, 12, 6),
        ("Heading 3", 12, BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_doc():
    md = SOURCE.read_text()
    lines = md.splitlines()
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    para_buf: list[str] = []
    table_buf: list[str] = []
    code_buf: list[str] = []
    in_code = False
    code_lang = ""
    body_started = False

    def flush_para():
        nonlocal para_buf
        if para_buf:
            text = " ".join(s.strip() for s in para_buf).strip()
            if text:
                if text.startswith("This section summarizes"):
                    add_callout(doc, "IT Review Focus", text, fill="FFF8E6")
                else:
                    add_para(doc, text)
            para_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_markdown_table(doc, parse_table(table_buf))
            table_buf = []

    def flush_code():
        nonlocal code_buf, code_lang
        if code_buf:
            text = "\n".join(code_buf)
            if code_lang == "mermaid":
                add_callout(doc, "Process Flow", "The end-user journey is represented below as a technical request sequence across the public domain, Kinsta routing, Resi WordPress, and the secure leasing portal handoff.", fill="EEF6FF")
                doc.add_page_break()
                add_process_flow(doc)
            else:
                add_code_block(doc, text)
            code_buf = []
            code_lang = ""

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                flush_para()
                flush_table()
                in_code = True
                code_lang = line.strip("`").strip()
            continue
        if in_code:
            code_buf.append(line)
            continue
        if line.startswith("|"):
            flush_para()
            table_buf.append(line)
            continue
        else:
            flush_table()
        if not line.strip():
            flush_para()
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            flush_para()
            body_started = True
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, line[3:].strip(), size=16, color=ACCENT)
            continue
        if not body_started and re.match(r"^(Status|Date|Audience|Scope):", line):
            continue
        if line.startswith("### "):
            flush_para()
            body_started = True
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, line[4:].strip(), size=13, color=ACCENT)
            continue
        if line.startswith("- "):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.167
            add_inline_runs(p, line[2:].strip())
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_para()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.167
            add_inline_runs(p, re.sub(r"^\d+\.\s+", "", line))
            continue
        para_buf.append(line)

    flush_para()
    flush_table()
    flush_code()

    doc.core_properties.title = "Venterra Web Platform Architecture"
    doc.core_properties.subject = "Kinsta-hosted public web platform framework"
    doc.core_properties.author = "Venterra Property Analytics"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
